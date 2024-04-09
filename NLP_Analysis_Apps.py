# -*- coding: utf-8 -*-
"""
Created on Wed Apr  3 18:51:50 2024

@author: vsingh1
"""

import streamlit as st
import pandas as pd
import spacy
from textblob import TextBlob
from wordcloud import WordCloud
import matplotlib.pyplot as plt
#import docx
from docx import Document  # Corrected import for python-docx
import pdfplumber
import string
import nltk
nltk.download('punkt')
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import re
from sklearn.feature_extraction.text import TfidfVectorizer
import altair as alt
import heapq


# Define the structure of your multi-page app
st.set_page_config(page_title="NLP Analysis Apps", layout="wide")

# Create a function for each app
def app1():
    # Load a pre-trained spaCy model for NLP tasks
    nlp = spacy.load("en_core_web_sm")

    # Define functions for reading PDF and DOCX files, preprocessing text, etc.
    def read_pdf(file):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + '\n'
        return text

    def read_docx(file):
        doc = Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + '\n'
        return text

    def preprocess_text(text):
        text = text.lower()
        doc = nlp(text)
        cleaned_tokens = [token.lemma_ for token in doc if not token.is_stop and not token.is_punct and not token.like_num]
        return " ".join(cleaned_tokens)

    def analyze_sentiment(text):
        blob = TextBlob(text)
        return blob.sentiment.polarity

    def extract_named_entities(text):
        doc = nlp(text)
        entities = set((ent.text, ent.label_) for ent in doc.ents if not ent.text.isspace())
        return list(entities)

    def extract_pos_tags(text):
        doc = nlp(text)
        pos_tags_dict = {}
        for token in doc:
            if not token.text.isspace() and not token.is_punct and not token.is_stop and len(token.text) > 1 and not any(char.isdigit() for char in token.text):
                if token.pos_ not in pos_tags_dict:
                    pos_tags_dict[token.pos_] = [token.text]
                else:
                    pos_tags_dict[token.pos_].append(token.text)
        pos_tags = [(pos, ", ".join(tokens)) for pos, tokens in pos_tags_dict.items()]
        return pos_tags

    def calculate_tfidf(text):
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform([text])
        df = pd.DataFrame(tfidf_matrix.toarray(), columns=vectorizer.get_feature_names_out())
        return df.T.sort_values(by=0, ascending=False)

    def generate_wordcloud(text):
        wordcloud = WordCloud(width=500, height=300, background_color='white').generate(text)
        fig, ax = plt.subplots(figsize=(10, 7))
        ax.imshow(wordcloud, interpolation='bilinear')
        ax.axis('off')
        st.pyplot(fig)

    # Streamlit app UI
    st.title("NLP Showcase")

    text_input_mode = st.radio("Choose input mode", ["Text Area", "Upload Document"])

    user_input = ""

    if text_input_mode == "Text Area":
        user_input = st.text_area("Enter text (up to 5000 characters)", value="", max_chars=5000, key="text_area", placeholder="Type or paste your text here...")
    else:
        uploaded_file = st.file_uploader("Upload a document", type=["txt", "pdf", "docx"], key="file_uploader")

        if uploaded_file is not None:
            if uploaded_file.type == "application/pdf":
                user_input = read_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                user_input = read_docx(uploaded_file)
            else:
                user_input = uploaded_file.getvalue().decode("utf-8")

    if st.button("Analyze") and user_input:
        if len(user_input.split()) > 5000:
            st.error("The text exceeds 5000 words. Please reduce the length.")
        else:
            with st.spinner('Processing text...'):
                cleaned_text = preprocess_text(user_input)
                sentiment = analyze_sentiment(cleaned_text)
                
                st.subheader("Sentiment Analysis")
                st.write("Sentiment: Positive" if sentiment > 0 else "Sentiment: Negative" if sentiment < 0 else "Sentiment: Neutral")

                st.subheader("Word Cloud")
                generate_wordcloud(cleaned_text)

                st.subheader("Word Frequencies")
                word_freq = pd.Series(cleaned_text.split()).value_counts().reset_index()
                word_freq.columns = ['Word', 'Frequency']
                freq_chart = alt.Chart(word_freq).mark_bar().encode(
                    x='Word',
                    y='Frequency',
                    color='Frequency',
                    tooltip=['Word', 'Frequency']
                ).interactive()
                st.altair_chart(freq_chart, use_container_width=True)

                entities = extract_named_entities(user_input)
                pos_tags = extract_pos_tags(user_input)
                tfidf_scores = calculate_tfidf(cleaned_text)

                with st.expander("Named Entities"):
                    st.dataframe(pd.DataFrame(entities, columns=["Entity", "Label"]))

                with st.expander("Part-of-Speech Tags"):
                    pos_tags_df = pd.DataFrame(pos_tags, columns=["POS Tag", "Tokens"])
                    st.table(pos_tags_df)

                with st.expander("TF-IDF Scores"):
                    st.dataframe(tfidf_scores.head(10))


def app2():
    # Load a pre-trained spaCy model for NLP tasks
    nlp = spacy.load("en_core_web_sm")

    # Define functions for reading PDF and DOCX files, preprocessing text, etc.
    def read_pdf(file):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + '\n'
        return text

    def read_docx(file):
        doc = Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + '\n'
        return text

    def preprocess_text(text):
        # Lowercase the text
        text = text.lower()
        
        # Remove numbers and special characters
        text = re.sub(r'\d+', '', text)
        text = text.translate(str.maketrans('', '', string.punctuation))
        
        # Tokenize the text
        tokens = word_tokenize(text)
        
        # Remove stop words
        stop_words = set(stopwords.words('english'))
        tokens = [word for word in tokens if word not in stop_words]
        
        # Lemmatization
        lemmatizer = WordNetLemmatizer()
        tokens = [lemmatizer.lemmatize(word) for word in tokens]
        
        # Join tokens back into text
        cleaned_text = ' '.join(tokens)
        
        return cleaned_text

    def analyze_sentiment(text):
        blob = TextBlob(text)
        return blob.sentiment.polarity

    def extract_named_entities(text):
        doc = nlp(text)
        entities = set((ent.text, ent.label_) for ent in doc.ents if not ent.text.isspace())
        return list(entities)

    def generate_wordcloud(text):
        wordcloud = WordCloud(width=500, height=300, background_color='white').generate(text)
        fig, ax = plt.subplots(figsize=(10, 7))
        ax.imshow(wordcloud, interpolation='bilinear')
        ax.axis('off')
        st.pyplot(fig)
        
    def summarize_text(text, n_sentences=3):
        # Ensure stop words are available within this function
        stop_words = set(stopwords.words('english'))

        # Tokenize the text into sentences
        doc = nlp(text)
        
        # Calculate word frequencies
        word_frequencies = {}
        for word in doc:
            if word.text.lower() not in stop_words:
                if word.text.lower() not in word_frequencies.keys():
                    word_frequencies[word.text.lower()] = 1
                else:
                    word_frequencies[word.text.lower()] += 1

        # Normalize frequencies
        maximum_frequency = max(word_frequencies.values())
        for word in word_frequencies.keys():
            word_frequencies[word] = (word_frequencies[word] / maximum_frequency)
        
        # Scoring sentences
        sentence_scores = {}
        for sent in doc.sents:
            for word in sent:
                if word.text.lower() in word_frequencies.keys():
                    if len(sent.text.split(' ')) < 30:
                        if sent not in sentence_scores.keys():
                            sentence_scores[sent] = word_frequencies[word.text.lower()]
                        else:
                            sentence_scores[sent] += word_frequencies[word.text.lower()]
        
        # Extract top n sentences as the summary
        summary_sentences = heapq.nlargest(n_sentences, sentence_scores, key=sentence_scores.get)
        summary = ' '.join([sent.text for sent in summary_sentences])

        return summary


    # Streamlit app UI
    st.title("Customer Feedback Analysis")

    text_input_mode = st.radio("Choose input mode", ["Text Area", "Upload Document"])

    user_input = ""

    if text_input_mode == "Text Area":
        user_input = st.text_area("Enter customer feedback (up to 5000 characters)", value="", max_chars=5000, key="text_area", placeholder="Type or paste customer feedback here...")
    else:
        uploaded_file = st.file_uploader("Upload a document", type=["txt", "pdf", "docx"], key="file_uploader")

        if uploaded_file is not None:
            if uploaded_file.type == "application/pdf":
                user_input = read_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                user_input = read_docx(uploaded_file)
            else:
                user_input = uploaded_file.getvalue().decode("utf-8")

    if st.button("Analyze") and user_input:
        if len(user_input.split()) > 5000:
            st.error("The text exceeds 5000 words. Please reduce the length.")
        else:
            with st.spinner('Processing customer feedback...'):
                cleaned_text = preprocess_text(user_input)
                sentiment = analyze_sentiment(cleaned_text)
                
                st.subheader("Sentiment Analysis")
                st.write("Sentiment: Positive" if sentiment > 0 else "Sentiment: Negative" if sentiment < 0 else "Sentiment: Neutral")

                st.subheader("Word Cloud")
                generate_wordcloud(cleaned_text)

                st.subheader("Feedback Summary")
                st.write("Total Words:", len(cleaned_text.split()))
        
                # Summarization
                summary = summarize_text(user_input)  # Using original text for summarization
                st.write("Summary:" , summary)


# Add pages to your app
pages = {
    "NLP Showcase": app1,
    "Customer Feedback Analysis": app2
}

st.sidebar.title('Apps')
selection = st.sidebar.radio("Go to", list(pages.keys()))
page = pages[selection]
page()
